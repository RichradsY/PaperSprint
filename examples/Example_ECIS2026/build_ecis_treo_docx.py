from pathlib import Path

from docx import Document


BASE = Path("/Users/ysx/Thesisplan/Manuscript/ICIS paper/Example_ECIS2026")
TEMPLATE = BASE / "ECIS2026_Submission_Template.docx"
OUTPUT = BASE / "ECIS_2026_Sprint3_TREO_Template_Ready.docx"

TITLE = "Productive but Pressured? Governing Enterprise GenAI Copilots Without Eroding Worker Dignity"
SHORT_TITLE = "GenAI Dignity Governance"
ABSTRACT = (
    "Enterprise generative AI copilots are spreading through professional service work, yet "
    "their consequences for workers remain undertheorized. This paper explains how organizational "
    "governance choices shape whether AI assistance improves productivity without degrading worker "
    "dignity. Drawing on affordance and algorithmic-control perspectives, it argues that dignity-"
    "preserving productivity depends on four conditions: transparency, overrideability, monitoring "
    "intensity, and task criticality. The paper defines dignity as accountable professional agency "
    "without degradation, excessive surveillance, or silenced contestation. It contributes a "
    "governance-based explanation of why the same copilot can function as an enabling partner in "
    "one setting and as algorithmic pressure in another. The paper also outlines a comparative "
    "qualitative roadmap for later empirical research in professional service firms."
)
KEYWORDS = "Keywords: Generative AI; worker dignity; algorithmic control; affordances; professional services"

SECTIONS = [
    (
        "Introduction",
        [
            "Enterprise generative AI copilots are spreading through professional service firms, where consultants, analysts, and advisory staff use them to draft, synthesize, and structure client-facing work. Early evidence suggests that generative AI can improve knowledge-worker productivity and output quality, but the gains are uneven and depend on task conditions and user capabilities (Dell'Acqua et al., 2023). Research on human-AI collaboration likewise shows that AI systems reshape rather than merely automate managerial and professional work (Sowa et al., 2021). What remains underdeveloped is a governance explanation of when these systems improve performance without degrading how workers are treated in the process.",
            "This paper addresses that gap by asking: How do enterprise GenAI copilot governance choices shape the balance between productivity and worker dignity in professional service knowledge work? The core claim is that the key issue is not simply whether organizations deploy copilots, but how they govern visibility, overrideability, monitoring, and accountability around them.",
        ],
    ),
    (
        "Theoretical Framing",
        [
            "An affordance perspective explains how GenAI copilots create action possibilities for faster drafting, broader information access, and more standardized work outputs (Leonardi, 2011). Yet affordances do not operate outside organizational control. Research on algorithmic control shows that data visibility, evaluation systems, and digitally mediated rules can intensify contested forms of control while redistributing discretion (Kellogg et al., 2020). Combining these lenses suggests that the consequences of enterprise GenAI depend less on the tool alone than on the governance regime surrounding its use.",
            "Worker dignity is the focal outcome because it captures a broader relational and normative condition than autonomy or well-being alone. Workplace dignity concerns whether workers are recognized as valued contributors rather than treated as replaceable instruments (Hodson, 2001; Lucas, 2015). Autonomy refers more narrowly to volition and self-direction in action (Gagne & Deci, 2005). Well-being concerns affective or health-related states. Control, by contrast, describes the governance arrangements that may shape these outcomes. In AI-mediated work, dignity is threatened when monitoring intensifies, discretion collapses, or blame is individualized while organizations still mandate AI use.",
        ],
    ),
    (
        "Emerging Propositions And Contribution",
        [
            "The paper argues that four governance conditions are especially consequential: transparency, overrideability, monitoring intensity, and task criticality. Together they shape whether the same copilot is experienced as enabling support or as algorithmic pressure.",
            "P1. Enterprise GenAI copilots are more likely to support dignity-preserving productivity when workers retain meaningful overrideability over AI outputs.",
            "P2. Monitoring-intensive copilot regimes are more likely to convert perceived AI affordances into experiences of algorithmic pressure.",
            "P3. The negative dignity effects of high monitoring and low overrideability are stronger in high-criticality, client-facing professional service tasks.",
            "P4. Governance regimes that combine transparency with shared accountability are more likely to sustain both productivity and worker dignity than regimes that individualize blame for AI-assisted failure.",
            "The paper contributes by reframing enterprise GenAI deployment as a governance problem rather than only an adoption or performance problem. It also gives conceptual precision to dignity as a worker-level outcome and identifies boundary conditions under which productivity gains remain compatible with responsible digital technology use.",
        ],
    ),
    (
        "Empirical Roadmap",
        [
            "Although this TREO submission is conceptual, it is designed to seed later empirical work. A plausible next step is a comparative qualitative case study of two to three professional service organizations, or business units within them, that have actively deployed GenAI copilots. Data could include semi-structured interviews, internal governance documents, training materials, and workflow artifacts. This roadmap matters because it shows how the conceptual argument can be translated into an empirical research program rather than remaining a thought piece.",
        ],
    ),
]

REFERENCES = [
    "Dell'Acqua, F., McFowland, E., Mollick, E. R., Lifshitz-Assaf, H., Kellogg, K., Rajendran, S., Krayer, L., Candelon, F., & Lakhani, K. R. (2023). Navigating the jagged technological frontier: Field experimental evidence of the effects of AI on knowledge worker productivity and quality. SSRN Electronic Journal. https://doi.org/10.2139/ssrn.4573321",
    "Gagne, M., & Deci, E. L. (2005). Self-determination theory and work motivation. Journal of Organizational Behavior, 26(4), 331-362. https://doi.org/10.1002/job.322",
    "Hodson, R. (2001). Dignity at work. Cambridge University Press. https://doi.org/10.1017/CBO9780511499333",
    "Kellogg, K. C., Valentine, M. A., & Christin, A. (2020). Algorithms at work: The new contested terrain of control. Academy of Management Annals, 14(1), 366-410. https://doi.org/10.5465/annals.2018.0174",
    "Leonardi, P. M. (2011). When flexible routines meet flexible technologies: Affordance, constraint, and the imbrication of human and material agencies. MIS Quarterly, 35(1), 147-167. https://doi.org/10.2307/23043493",
    "Lucas, K. (2015). Workplace dignity: Communicating inherent, earned, and remediated dignity. Journal of Management Studies, 52(5), 621-646. https://doi.org/10.1111/joms.12133",
    "Sowa, K., Przegalinska, A., & Ciechanowski, L. (2021). Cobots in knowledge work. Journal of Business Research, 125, 135-142. https://doi.org/10.1016/j.jbusres.2020.11.038",
]


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def delete_table(table):
    element = table._element
    parent = element.getparent()
    parent.remove(element)


doc = Document(str(TEMPLATE))
doc.core_properties.author = ""
doc.core_properties.last_modified_by = ""
doc.core_properties.comments = ""
doc.core_properties.title = TITLE

doc.sections[0].header.paragraphs[0].text = SHORT_TITLE
doc.sections[0].header.paragraphs[0].style = doc.styles["Header"]

doc.paragraphs[0].text = TITLE
doc.paragraphs[1].text = "TREO Paper"
doc.paragraphs[9].text = "Abstract"
doc.paragraphs[10].text = ABSTRACT
doc.paragraphs[12].text = KEYWORDS

keep = {0, 1, 2, 3, 9, 10, 11, 12}
for idx, paragraph in reversed(list(enumerate(list(doc.paragraphs)))):
    if idx not in keep:
        delete_paragraph(paragraph)

for table in list(doc.tables):
    delete_table(table)

for heading, paragraphs in SECTIONS:
    doc.add_paragraph(heading, style="Heading 1")
    for text in paragraphs:
        doc.add_paragraph(text, style="Basic text")

doc.add_paragraph("References", style="Subtitle")
for ref in REFERENCES:
    doc.add_paragraph(ref, style="Reference")

doc.save(str(OUTPUT))
print(f"Wrote {OUTPUT}")
